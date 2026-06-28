import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('shd'):
            tcPr.remove(child)
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    tcPr.append(parse_xml(shading_xml))

def set_cell_borders(cell, top="none", bottom="none", left="none", right="none", 
                      top_color="auto", bottom_color="auto", left_color="auto", right_color="auto",
                      top_sz="0", bottom_sz="0", left_sz="0", right_sz="0"):
    tcPr = cell._tc.get_or_add_tcPr()
    for child in list(tcPr):
        if child.tag.endswith('tcBorders'):
            tcPr.remove(child)
    borders_xml = f'''<w:tcBorders {nsdecls("w")}>
        <w:top w:val="{top}" w:sz="{top_sz}" w:space="0" w:color="{top_color}"/>
        <w:left w:val="{left}" w:sz="{left_sz}" w:space="0" w:color="{left_color}"/>
        <w:bottom w:val="{bottom}" w:sz="{bottom_sz}" w:space="0" w:color="{bottom_color}"/>
        <w:right w:val="{right}" w:sz="{right_sz}" w:space="0" w:color="{right_color}"/>
    </w:tcBorders>'''
    tcPr.append(parse_xml(borders_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clean_equation_text(text):
    """
    Cleans LaTeX formatting from markdown text to render readable, professional plain-text equations in Word.
    """
    cleaned = text.strip()
    if cleaned.startswith('$$') and cleaned.endswith('$$'):
        cleaned = cleaned[2:-2].strip()
    elif cleaned.startswith('$') and cleaned.endswith('$'):
        cleaned = cleaned[1:-1].strip()
        
    replacements = [
        (r'\mathbf{x}', 'x'),
        (r'\mathbf{w}', 'w'),
        (r'\mathbf{X}', 'X'),
        (r'\mathbf{W}', 'W'),
        (r'\mathbf{Y}', 'Y'),
        (r'\mathbb{R}', 'ℝ'),
        (r'\times', ' × '),
        (r'\sum', '∑'),
        (r'\sigma', 'σ'),
        (r'\alpha', 'α'),
        (r'\beta', 'β'),
        (r'\epsilon', 'ε'),
        (r'\delta', 'δ'),
        (r'\theta', 'θ'),
        (r'\hat{y}', 'y_hat'),
        (r'\text{T}', 'T'),
        (r'\phi', 'φ'),
        (r'\gamma', 'γ'),
        (r'\eta', 'η'),
        (r'\mu', 'μ'),
        (r'\partial', '∂'),
        (r'\nabla', '∇'),
        (r'\approx', ' ≈ '),
        (r'\neq', ' ≠ '),
        (r'\leq', ' ≤ '),
        (r'\geq', ' ≥ '),
        (r'\infty', '∞'),
        (r'\dots', '...'),
    ]
    
    for pattern, repl in replacements:
        cleaned = cleaned.replace(pattern, repl)
        
    cleaned = re.sub(r'\\(hat|tilde|bar|vec)\{(.*?)\}', r'\2_hat', cleaned)
    cleaned = re.sub(r'\_\{(.*?)\}', r'_\1', cleaned)
    cleaned = re.sub(r'\^\{(.*?)\}', r'^\1', cleaned)
    cleaned = re.sub(r'\\(mathrm|mathbf|mathit|mathsf|mathbb)\{(.*?)\}', r'\2', cleaned)
    cleaned = re.sub(r'\\(text|textrm)\{(.*?)\}', r'\2', cleaned)
    cleaned = re.sub(r'\\left\(', '(', cleaned)
    cleaned = re.sub(r'\\right\)', ')', cleaned)
    cleaned = re.sub(r'\\left\[', '[', cleaned)
    cleaned = re.sub(r'\\right\]', ']', cleaned)
    cleaned = re.sub(r'\\([a-zA-Z]+)', r'\1', cleaned) # strip any lingering backslashes
    return cleaned.strip()

def parse_inline_formatting(paragraph, text, base_font_size=12, italic_all=False):
    """
    Parses **bold** and *italic* markdown syntax and adds styled runs to the paragraph.
    """
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*)')
    parts = pattern.split(text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
            if italic_all:
                run.italic = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(base_font_size)
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(base_font_size)
                if italic_all:
                    run.italic = True

def markdown_to_docx(md_path, docx_path):
    print(f"[Compiler] Compiling {md_path} -> {docx_path}...")
    doc = Document()
    
    # Configure page margins (1 inch / 72pt = 1440 dxa)
    sections = doc.sections
    for section in sections:
        section.top_margin = Pt(72)
        section.bottom_margin = Pt(72)
        section.left_margin = Pt(72)
        section.right_margin = Pt(72)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_list = False
    in_table = False
    
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # 0. Parse Code Blocks (using a beautiful 1x1 table for consistent shading and borders)
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if next_line.strip().startswith('```'):
                    break
                code_lines.append(next_line)
                i += 1
            
            if code_lines:
                table = doc.add_table(rows=1, cols=1)
                table.style = 'Normal Table'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = False
                
                # Full page text width in Word (8.5 inch width - 2 inch margins = 6.5 inches)
                # 6.5 inches = 468 points = 9360 dxa
                table.columns[0].width = Pt(468)
                cell = table.cell(0, 0)
                cell.width = Pt(468)
                
                set_cell_background(cell, "F5F5F5") # Light grey background
                set_cell_borders(cell, left="single", left_color="22703F", left_sz="24") # Solid thick Forest Green left border (3pt)
                set_cell_margins(cell, top=140, bottom=140, left=200, right=200) # Premium inner padding
                
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                
                for idx, c_line in enumerate(code_lines):
                    text_to_add = c_line.rstrip('\r\n')
                    if idx < len(code_lines) - 1:
                        text_to_add += '\n'
                    run = p.add_run(text_to_add)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9.5)
                    run.font.color.rgb = RGBColor(33, 33, 33) # Charcoal code text
            
            i += 1
            continue
            
        # Skip empty lines, but break list grouping
        if not stripped:
            in_list = False
            i += 1
            continue
            
        # 1. Parse Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)$', stripped)
        if heading_match:
            in_list = False
            level = len(heading_match.group(1))
            title_text = heading_match.group(2)
            
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.keep_with_next = True
            
            if level == 1:
                run = p.add_run(title_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(34, 112, 63) # Deep Forest Green accent for FLORA
                p.paragraph_format.space_before = Pt(24)
            elif level == 2:
                run = p.add_run(title_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(46, 125, 50)
            elif level == 3:
                run = p.add_run(title_text)
                run.bold = True
                run.italic = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(66, 66, 66)
            else:
                run = p.add_run(title_text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(97, 97, 97)
                
            i += 1
            continue

        # 2. Parse Visual Placeholders [INSERT IMAGE HERE: ...]
        if stripped.startswith('[INSERT IMAGE HERE:'):
            in_list = False
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            placeholder_text = stripped[1:-1]
            run = p.add_run(f"\n📎 {placeholder_text}\n")
            run.bold = True
            run.italic = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(21, 101, 192) # Soft academic blue
            
            i += 1
            continue

        # 3. Parse Bullet Points / Lists
        list_match = re.match(r'^[\-\*\u2022]\s+(.*)$', stripped)
        if list_match:
            content = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, content, base_font_size=12)
            in_list = True
            i += 1
            continue

        # 4. Parse Numbered Lists
        num_list_match = re.match(r'^(\d+)\.\s+(.*)$', stripped)
        if num_list_match:
            num = num_list_match.group(1)
            content = num_list_match.group(2)
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            parse_inline_formatting(p, content, base_font_size=12)
            in_list = True
            i += 1
            continue

        # 5. Parse Table
        if stripped.startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i].strip())
                i += 1
            
            table_rows = [r for r in table_rows if not re.match(r'^\|[\s\-:|]+$', r)]
            
            if len(table_rows) > 0:
                parsed_rows = [[cell.strip() for cell in row.split('|')[1:-1]] for row in table_rows]
                num_cols = max(len(r) for r in parsed_rows)
                
                table = doc.add_table(rows=len(parsed_rows), cols=num_cols)
                table.style = 'Normal Table'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                table.autofit = True
                
                for r_idx, row_data in enumerate(parsed_rows):
                    row = table.rows[r_idx]
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < len(row.cells):
                            cell = row.cells[c_idx]
                            cell.text = ""
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            p = cell.paragraphs[0]
                            p.paragraph_format.space_after = Pt(4)
                            p.paragraph_format.space_before = Pt(4)
                            
                            val_clean = cell_value.strip()
                            is_short_or_num = re.match(r'^[\d\.\%\-\/]+$', val_clean) or len(val_clean) < 10
                            
                            if r_idx == 0:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = p.add_run(cell_value)
                                run.bold = True
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
                                run.font.color.rgb = RGBColor(255, 255, 255) # White text
                                set_cell_background(cell, "22703F") # Deep Green header
                                set_cell_borders(cell, 
                                                 top="single", top_color="22703F", top_sz="4",
                                                 bottom="single", bottom_color="1B5E20", bottom_sz="12",
                                                 left="single", left_color="22703F", left_sz="4",
                                                 right="single", right_color="22703F", right_sz="4")
                            else:
                                if is_short_or_num:
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                                
                                parse_inline_formatting(p, cell_value, base_font_size=9.5)
                                
                                # Zebra striping
                                if r_idx % 2 == 0:
                                    set_cell_background(cell, "F4F9F4") # Light Green tint
                                else:
                                    set_cell_background(cell, "FFFFFF")
                                    
                                set_cell_borders(cell, 
                                                 top="single", top_color="E0E0E0", top_sz="4",
                                                 bottom="single", bottom_color="E0E0E0", bottom_sz="4",
                                                 left="single", left_color="E0E0E0", left_sz="4",
                                                 right="single", right_color="E0E0E0", right_sz="4")
                                
                            set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                
                doc.add_paragraph()
            continue

        # 5.5 Detect and clean Standalone Equations
        is_equation = False
        if (not stripped.startswith('#') and
            not stripped.startswith('-') and
            not stripped.startswith('*') and
            not stripped.startswith('|') and
            not stripped.startswith('[') and
            '=' in stripped):
            
            words_in_line = stripped.split()
            math_keywords = {'sum', 'from', 'to', 'max', 'min', 'if', 'and', 'constant', 'otherwise', 'where', 'or', 'for', 'exp', 'log', 'floor', 'ceil', 'sqrt', 's.t.'}
            non_math_words = [w for w in words_in_line if w.lower().rstrip(',.:;') not in math_keywords and w.isalnum() and not any(char.isdigit() for char in w)]
            
            if len(non_math_words) <= 6:
                is_equation = True
            elif any(op in stripped for op in ['+', '∑', '^', '_', '\\', 'prod', 'beta', 'alpha', 'gamma', 'lambda', 'theta', 'epsilon', 'delta']):
                is_equation = True

        if is_equation:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            cleaned_eq = clean_equation_text(stripped)
            run = p.add_run(cleaned_eq)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.italic = True
            run.font.color.rgb = RGBColor(33, 33, 33)
            
            i += 1
            continue

        # 6. Parse regular paragraph
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.15
        
        parse_inline_formatting(p, stripped, base_font_size=12)
        i += 1

    try:
        doc.save(docx_path)
        print(f"[Compiler] Successfully generated {docx_path}")
    except PermissionError:
        print(f"\n[Warning] Permission Denied: Could not write to {docx_path}")
        print("          This usually means the file is open in Microsoft Word.")
        print("          Please close the file in Word and run the compiler again.\n")
    except Exception as e:
        print(f"[Error] Failed to save {docx_path}: {e}")

if __name__ == "__main__":
    import sys
    doc_dir = os.path.dirname(os.path.abspath(__file__))
    
    if len(sys.argv) > 1:
        target = sys.argv[1]
        if not os.path.isabs(target):
            target = os.path.join(doc_dir, target)
        
        if os.path.exists(target) and target.endswith('.md'):
            out_docx = target.replace('.md', '.docx')
            markdown_to_docx(target, out_docx)
        else:
            print(f"[Error] File not found or not markdown: {target}")
    else:
        print("[Compiler] Running in bulk compilation mode...")
        compiled_count = 0
        for f_name in os.listdir(doc_dir):
            if f_name.endswith(".md") and f_name != "flora_thesis_blueprint.md":
                md_path = os.path.join(doc_dir, f_name)
                docx_path = os.path.join(doc_dir, f_name.replace(".md", ".docx"))
                markdown_to_docx(md_path, docx_path)
                compiled_count += 1
        print(f"[Compiler] Bulk compilation complete. Compiled {compiled_count} files.")
